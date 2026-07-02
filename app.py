from flask import Flask, render_template, request, send_file, redirect, url_for
import os, json
from generate_text import generate_text
from docx import Document
from docx.shared import Pt
from fpdf import FPDF

app = Flask(__name__)

UPLOAD_FOLDER = "static/uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER

# This will store the last generated resume with AI layout info
last_resume = {}


# ------------------ HOME ------------------
@app.route("/")
def home():
    return redirect(url_for("resume"))


# ------------------ RESUME PAGE ------------------
@app.route("/resume", methods=["GET", "POST"])
def resume():
    global last_resume

    student = None
    summary = None
    template = "minimal"
    photo = None

    if request.method == "POST":
        profession = request.form.get("profession", "")
        name = request.form.get("name", "")
        skills = request.form.get("skills", "")
        career = request.form.get("career", "")
        template = request.form.get("template", "minimal")

        # Save photo
        file = request.files.get("photo")
        if file and file.filename:
            filename = file.filename.replace(" ", "_")
            filepath = os.path.join(app.config["UPLOAD_FOLDER"], filename)
            file.save(filepath)
            photo = "/" + filepath.replace("\\", "/")

        student = {
            "profession": profession.title(),
            "name": name.title(),
            "skills": [s.strip() for s in skills.split(",") if s.strip()],
            "career": career
        }

        # ---------- AI PROMPT (TEXT + LAYOUT) ----------
        prompt = f"""
You are an ATS resume designer.

Return ONLY valid JSON in this format:

{{
  "summary": "Professional resume summary here",
  "layout": {{
    "font_size": 11,
    "line_spacing": 1.5,
    "section_spacing": 12,
    "page_margin": 20
  }}
}}

Rules:
- Summary must be 2–3 sentences.
- Layout must be ATS safe.
- No markdown, no explanation, only JSON.

Data:
Profession: {profession}
Name: {name}
Skills: {skills}
Career Goal: {career}
"""

        ai_response = generate_text(prompt)

        # Parse AI response
        try:
            ai_data = json.loads(ai_response)
            summary = ai_data["summary"]
            layout = ai_data["layout"]
        except:
            # Fallback if AI response is not valid JSON
            summary = ai_response
            layout = {
                "font_size": 11,
                "line_spacing": 1.5,
                "section_spacing": 12,
                "page_margin": 20
            }

        # Store everything for PDF / DOCX
        last_resume = {
            "name": student["name"],
            "profession": student["profession"],
            "skills": ", ".join(student["skills"]),
            "career": student["career"],
            "summary": summary,
            "layout": layout,
            "template": template,
            "photo": photo
        }

    return render_template(
        "resume.html",
        student=student,
        summary=summary,
        template=template,
        photo=photo
    )


# ------------------ DOWNLOAD PDF (AI FORMATTED) ------------------
@app.route("/download/pdf", methods=["POST"])
def download_pdf():
    if not last_resume:
        return "Generate a resume first."

    layout = last_resume["layout"]

    content = f"""
{last_resume['name']}
{last_resume['profession']}

Career Goal:
{last_resume['career']}

Skills:
{last_resume['skills']}

Professional Summary:
{last_resume['summary']}
"""

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=layout["page_margin"])
    pdf.add_page()
    pdf.set_font("Arial", size=layout["font_size"])
    pdf.multi_cell(0, layout["line_spacing"] * 8, content)

    file_path = "resume_ai_formatted.pdf"
    pdf.output(file_path)

    return send_file(file_path, as_attachment=True)


# ------------------ DOWNLOAD DOCX (AI FORMATTED) ------------------
@app.route("/download/docx", methods=["POST"])
def download_docx():
    if not last_resume:
        return "Generate a resume first."

    layout = last_resume["layout"]

    doc = Document()

    # Set base font size
    style = doc.styles["Normal"]
    style.font.size = Pt(layout["font_size"])

    # Header
    doc.add_heading(last_resume["name"], level=1)
    doc.add_paragraph(last_resume["profession"])

    # Career Goal
    p = doc.add_paragraph(f"Career Goal:\n{last_resume['career']}")
    p.paragraph_format.space_after = Pt(layout["section_spacing"])

    # Skills
    p = doc.add_paragraph(f"Skills:\n{last_resume['skills']}")
    p.paragraph_format.space_after = Pt(layout["section_spacing"])

    # Summary
    doc.add_heading("Professional Summary", level=2)
    p = doc.add_paragraph(last_resume["summary"])
    p.paragraph_format.space_after = Pt(layout["section_spacing"])

    file_path = "resume_ai_formatted.docx"
    doc.save(file_path)

    return send_file(file_path, as_attachment=True)


# ------------------ PORTFOLIO PAGE ------------------
@app.route("/portfolio", methods=["GET", "POST"])
def portfolio():
    portfolio_data = None

    if request.method == "POST":
        portfolio_data = {
            "name": request.form.get("name", ""),
            "title": request.form.get("title", ""),
            "about": request.form.get("about", ""),
            "skills": request.form.get("skills", ""),
            "projects": request.form.get("projects", ""),
            "contact": request.form.get("contact", "")
        }

    return render_template("portfolio.html", data=portfolio_data)


# ------------------ RUN APP ------------------
if __name__ == "__main__":
    app.run(debug=True)
